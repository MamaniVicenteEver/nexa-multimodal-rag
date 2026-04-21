"""
Extractor unificado de texto para múltiples formatos de archivo.
Soporta: PDF, DOCX, TXT, MD, JSON, e imágenes (vía OCR).
"""

import os
import json
import io
from typing import List, Union

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from src.core.domain.extracted_page import ExtractedPage
from src.core.ports.ocr_provider import IOCRProvider
from src.infrastructure.ocr.processors.factory import get_page_processor
from src.modules.ingestion.hybrid_router import HybridDocumentProcessor
from src.shared.logging import get_logger

logger = get_logger("document_extractor")


class DocumentExtractor:
    def __init__(self, ocr_adapter: IOCRProvider):
        self.ocr_adapter = ocr_adapter
        self._hybrid_processor = None

    async def _get_hybrid_processor(self):
        if self._hybrid_processor is None:
            from src.infrastructure.ocr.processors.factory import get_page_processor
            from src.modules.ingestion.hybrid_router import HybridDocumentProcessor
            page_processor = get_page_processor()
            self._hybrid_processor = HybridDocumentProcessor(page_processor=page_processor)
        return self._hybrid_processor

    async def extract(self, file_bytes: bytes, mime_type: str, filename: str, doc_id: str) -> List[ExtractedPage]:
        _, ext = os.path.splitext(filename)
        ext = ext.lower()

        if mime_type == "application/pdf" or ext == ".pdf":
            logger.info(f"Extrayendo PDF: {filename}")
            hybrid_proc = await self._get_hybrid_processor()
            return await hybrid_proc.process_pdf(file_bytes, doc_id)

        # ---- DOCX (Word) ----
        elif mime_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword"
        ) or ext in (".docx", ".doc"):
            logger.info(f"Extrayendo DOCX: {filename}")
            text = self._extract_docx(file_bytes)
            return [self._make_single_page(text, doc_id, "DOCX")]

        # ---- Imágenes (vía OCR directo) ----
        elif mime_type.startswith("image/"):
            logger.info(f"Extrayendo imagen con OCR: {filename}")
            text = await self.ocr_adapter.extract(file_bytes)
            return [self._make_single_page(text, doc_id, "IMAGE")]

        # ---- Texto plano, Markdown, JSON ----
        elif mime_type.startswith("text/") or ext in (".txt", ".md", ".markdown", ".json"):
            logger.info(f"Extrayendo texto/{ext} directamente: {filename}")
            text = self._extract_text_or_json(file_bytes, ext)
            return [self._make_single_page(text, doc_id, "TEXT")]

        else:
            raise ValueError(f"Formato no soportado: {mime_type} / {ext}")

    def _make_single_page(self, content: str, doc_id: str, page_type: str) -> ExtractedPage:
        """Crea una única ExtractedPage para formatos no paginados."""
        return ExtractedPage(
            page_number=1,
            content=content,
            page_type=page_type,
            metadata={"doc_id": doc_id, "page_index": 0}
        )

    def _extract_docx(self, file_bytes: bytes) -> str:
        """
        Extrae texto de un documento Word (.docx) y lo convierte a Markdown básico.
        """
        doc = DocxDocument(io.BytesIO(file_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # Detectar estilo para encabezados
            if para.style.name and para.style.name.startswith('Heading'):
                level = 1
                if para.style.name[-1].isdigit():
                    level = int(para.style.name[-1])
                paragraphs.append(f"{'#' * level} {text}")
            else:
                paragraphs.append(text)
        # También extraer texto de tablas (simplificado)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n\n".join(paragraphs)

    def _extract_text_or_json(self, file_bytes: bytes, ext: str) -> str:
        """
        Lee archivos de texto (.txt, .md) o devuelve el contenido crudo de JSON.
        """
        content = file_bytes.decode('utf-8', errors='ignore')
        if ext == ".json":
            return content
        else:
            return content